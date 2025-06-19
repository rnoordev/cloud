# Cloud-Based Document Analytics Program with Flask Web Interface
# Features: Upload, Sort, Search, Classify PDF/Word documents stored on the cloud

import os
import time
import fitz  # PyMuPDF
import docx
from flask import Flask, request, render_template, redirect, url_for, send_from_directory, flash
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.pipeline import make_pipeline
import joblib

app = Flask(__name__)
app.secret_key = 'secret123'
UPLOAD_FOLDER = "cloud_storage"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# === Utility Functions ===
def upload_document(file):
    filename = file.filename
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(path)
    return filename

def extract_title(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        doc = fitz.open(file_path)
        title = doc[0].get_text().split("\n")[0].strip()
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        title = doc.paragraphs[0].text.strip()
    else:
        title = os.path.basename(file_path)
    return title

def extract_text(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        doc = fitz.open(file_path)
        return "\n".join([page.get_text() for page in doc])
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""

def sort_documents_by_title():
    files = os.listdir(UPLOAD_FOLDER)
    files_with_titles = [(extract_title(os.path.join(UPLOAD_FOLDER, f)), f) for f in files]
    files_sorted = sorted(files_with_titles, key=lambda x: x[0])
    return [f[1] for f in files_sorted]

def search_documents(keywords):
    matches = []
    for filename in os.listdir(UPLOAD_FOLDER):
        path = os.path.join(UPLOAD_FOLDER, filename)
        text = extract_text(path).lower()
        if all(keyword.lower() in text for keyword in keywords):
            matches.append(filename)
    return matches

def train_classifier(documents, labels):
    texts = [extract_text(os.path.join(UPLOAD_FOLDER, f)) for f in documents]
    model = make_pipeline(TfidfVectorizer(), MultinomialNB())
    model.fit(texts, labels)
    joblib.dump(model, "classifier.joblib")

def classify_documents():
    if not os.path.exists("classifier.joblib"):
        return []
    model = joblib.load("classifier.joblib")
    results = []
    for f in os.listdir(UPLOAD_FOLDER):
        text = extract_text(os.path.join(UPLOAD_FOLDER, f))
        label = model.predict([text])[0]
        results.append((f, label))
    return results

def get_statistics():
    num_files = len(os.listdir(UPLOAD_FOLDER))
    total_size = sum(os.path.getsize(os.path.join(UPLOAD_FOLDER, f)) for f in os.listdir(UPLOAD_FOLDER))
    return {"count": num_files, "size_kb": total_size / 1024}

# === Flask Routes ===
@app.route("/")
def index():
    stats = get_statistics()
    return render_template("index.html", stats=stats)

@app.route("/upload", methods=["POST"])
def upload():
    file = request.files['file']
    try:
        upload_document(file)
        flash("✅ File uploaded successfully.")
    except Exception as e:
        flash(f"⚠️ Upload error: {str(e)}")
    return redirect(url_for('index'))


@app.route("/search", methods=["POST"])
def search():
    keywords = request.form['keywords'].split(',')
    results = search_documents(keywords)
    return render_template("search.html", results=results)

@app.route("/sort")
def sort():
    sorted_files = sort_documents_by_title()
    return render_template("sort.html", files=sorted_files)

@app.route("/classify")
def classify():
    results = classify_documents()
    if not results:
        flash("Model not trained yet. Please train the classifier first.")
        return redirect(url_for('train'))
    return render_template("classify.html", results=results)

@app.route("/train", methods=["GET", "POST"])
def train():
    if request.method == "POST":
        files = request.form.getlist("file")
        labels = request.form.getlist("label")
        if len(files) == len(labels) and files:
            train_classifier(files, labels)
            flash("Model trained successfully.")
            return redirect(url_for('index'))
        else:
            flash("Error: Please make sure each file has a corresponding label.")
    file_list = os.listdir(UPLOAD_FOLDER)
    return render_template("train.html", files=file_list)

@app.route("/download/<filename>")
def download(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True)
